"""Generate Uniware Project Status .docx for Niki."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path(__file__).resolve().parent.parent / "Documentations" / "Uniware_Project_Status_For_Niki.docx"


def para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(11)
        r.font.name = "Calibri"


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Uniware Website — Project Status")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("For Niki · Prepared by Srimathi · 2 September 2026")
    sr.font.size = Pt(10)
    sr.italic = True
    sr.font.name = "Calibri"
    doc.add_paragraph()

    doc.add_heading("1. Executive summary", level=1)
    para(
        doc,
        "New site development continues on global.uniware.net. Significant work is complete in code "
        "locally but not yet pushed to staging. Dhana confirmed: do not change uniware.net DNS until "
        "the entire site is built and validated. We need Niki to confirm this is the official cutover approach.",
    )

    doc.add_heading("2. Done (local — not all pushed)", level=1)
    doc.add_heading("Case study CMS (additive — safe for published studies)", level=2)
    for x in [
        "Problem & Solution body: bold, bullet/numbered lists, sub-heading (24px Space Grotesk)",
        "Updated Problem body helper text",
        "Optional Additional Section (toggle, after Quote after Problem)",
        "Frontend rendering + CSS",
    ]:
        bullet(doc, x)

    doc.add_heading("Cutover infrastructure", level=2)
    for x in [
        "~50+ WordPress 301 redirects (cutover-redirects.ts)",
        "WordPress fallback for unbuilt pages (next.config.ts + env var)",
        "SentinelOne redirect: /sentinel-one-cloud-security/ → Cloud Security",
        "GA4 component ready (needs measurement ID)",
        "Go-live plan doc for Dhana (separate attachment)",
    ]:
        bullet(doc, x)

    doc.add_heading("Mux / Studio", level=2)
    bullet(doc, "Video asset list shows filename instead of 'ready'")
    bullet(doc, "Unused dummy Mux asset removed")

    doc.add_heading("Already live on staging/main (earlier work)", level=2)
    bullet(doc, "Homepage, cloud pages, AWS hub, RDS, GenAI, booking panel, AI Solutions nav, Customers CMS, case study Mux fix")

    doc.add_heading("3. Pending", level=1)
    table(
        doc,
        ["Item", "Owner", "Priority"],
        [
            ["Push local changes to staging", "Srimathi", "High"],
            ["Test on global.uniware.net", "Srimathi", "High"],
            ["Deploy hosted Sanity Studio", "Srimathi", "Medium"],
            ["Set WORDPRESS_FALLBACK_ORIGIN on Vercel", "Srimathi", "High"],
            ["GA4 baseline from old WordPress", "Srimathi", "Medium"],
            ["PDF export for AWS Partner Central", "Srimathi", "Plan only — not built"],
            ["Final uniware.net DNS cutover", "Dhana + Niki", "Deferred"],
            ["AWS WordPress IP for post-cutover fallback", "Dhana", "Deferred"],
        ],
    )

    doc.add_heading("4. What is blocking", level=1)
    table(
        doc,
        ["Blocker", "Who", "Impact"],
        [
            ["Confirm cutover approach (incremental vs deferred)", "Niki", "Pauses DNS go-live requests to Dhana"],
            ["Push to staging not done yet", "Srimathi", "New work not on global.uniware.net"],
            ["Hosted Studio not redeployed", "Srimathi", "Editors don't see new case study fields"],
            ["Full site + sign-off before final cutover", "All", "Expected — uniware.net stays on WordPress"],
        ],
    )

    doc.add_heading("5. Current approach (per Dhana)", level=1)
    para(doc, "uniware.net → old WordPress (production unchanged)")
    para(doc, "global.uniware.net → new Vercel site (development & testing)")
    para(doc, "Final cutover → when entire site is complete + Niki sign-off")

    doc.add_heading("6. Decisions needed from Niki", level=1)
    for x in [
        "Cutover: incremental (cutover doc) or deferred until full site (Dhana)?",
        "Approve push to staging for case study + cutover work?",
        "PDF export: review plan when back from holiday?",
        "Confirm /aws-system-manager/ → RDS redirect?",
    ]:
        bullet(doc, x)

    doc.add_heading("7. Recommended next steps", level=1)
    for i, x in enumerate(
        [
            "Niki confirms cutover approach",
            "Srimathi pushes to staging + sets Vercel env vars",
            "Test redirects + fallback on global.uniware.net",
            "Deploy Sanity Studio",
            "Validate case study CMS on draft study",
            "Later: GA4, PDF build, final cutover with Dhana",
        ],
        1,
    ):
        bullet(doc, f"{i}. {x}")

    doc.add_paragraph()
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = f.add_run("Uniware Systems · September 2026")
    fr.font.size = Pt(9)
    fr.italic = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()

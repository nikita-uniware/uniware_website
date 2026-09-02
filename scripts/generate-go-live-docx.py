"""Generate Uniware.net Go-Live Plan .docx for email attachment."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path(__file__).resolve().parent.parent / "Documentations" / "Uniware_net_Go_Live_Plan.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("uniware.net Website Go-Live Plan")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Prepared for: Dhana (DNS / Infrastructure) · Niki (Sign-off)\nPrepared by: Srimathi · Uniware Website Rebuild")
    sr.font.size = Pt(10)
    sr.font.name = "Calibri"
    sr.italic = True

    doc.add_paragraph()

    add_heading(doc, "1. Summary", level=1)
    add_para(
        doc,
        "We are moving uniware.net from the old WordPress site (AWS) to the new Next.js site on Vercel. "
        "Rebuilt pages are served directly from Vercel. Pages not yet rebuilt are fetched from the old "
        "WordPress server automatically, so nothing breaks during the transition.",
    )
    add_para(
        doc,
        "global.uniware.net already points to Vercel and is used for testing. Go-live means updating "
        "DNS in Bluehost so uniware.net and www.uniware.net also point to Vercel.",
    )

    add_heading(doc, "2. Email / DNS — Important Clarification", level=1)
    add_para(
        doc,
        "Changing the website A record does NOT affect email, as long as MX, SPF, and DKIM records "
        "remain unchanged. Email and website use separate DNS records.",
        bold=True,
    )
    add_table(
        doc,
        ["Record type", "Name", "Purpose", "Action at go-live"],
        [
            ["MX", "—", "Email delivery (sales@uniware.net, etc.)", "DO NOT CHANGE"],
            ["SPF / DKIM", "—", "Email authentication", "DO NOT CHANGE"],
            ["A", "@ (root)", "Website at uniware.net", "Change to Vercel IP (see Section 5)"],
            ["CNAME", "www", "Website at www.uniware.net", "Change to Vercel CNAME (see Section 5)"],
        ],
    )

    add_heading(doc, "3. What Is Already Done (Development)", level=1)
    for item in [
        "301 redirects for all rebuilt pages (old WordPress URLs → new site URLs)",
        "WordPress fallback rule for pages not yet rebuilt (e.g. /about-us/, /career/)",
        "Google Analytics (GA4) ready — activates when measurement ID is set",
        "Testing domain global.uniware.net configured on Vercel",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. Pre Go-Live Checklist", level=1)
    add_heading(doc, "4.1 Srimathi (before asking for DNS change)", level=2)
    for item in [
        "Push latest code to staging and deploy on Vercel",
        "Set Vercel env: WORDPRESS_FALLBACK_ORIGIN = https://uniware.net (for pre-cutover testing)",
        "Test on global.uniware.net: new pages, redirects, and fallback pages",
        "Obtain Niki sign-off that testing is complete",
        "Add uniware.net (apex) in Vercel Domains and note exact A record IP shown",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.2 Dhana (before DNS change)", level=2)
    for item in [
        "Provide direct AWS IP or hostname of the old WordPress server (needed after go-live for fallback)",
        "Export current DNS records from Bluehost — confirm MX, SPF, DKIM are documented",
        "Confirm no other services depend on the current @ A record besides the website",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.3 Niki", level=2)
    add_bullet(doc, "Review test results on global.uniware.net and approve DNS go-live date")

    add_heading(doc, "5. DNS Changes in Bluehost (Go-Live Day)", level=1)
    add_para(doc, "Only perform these changes after Niki confirms testing is complete.", bold=True)
    add_para(doc, "Recommended order:")

    add_heading(doc, "Step 1 — www subdomain (lower risk, test first)", level=2)
    add_table(
        doc,
        ["Type", "Name / Host", "Value"],
        [
            ["CNAME", "www", "66fd6620af51ddf4.vercel-dns-017.com"],
        ],
    )
    add_para(doc, "Test: open https://www.uniware.net — should show the new Vercel site.")

    add_heading(doc, "Step 2 — Root domain (main cutover)", level=2)
    add_para(
        doc,
        "Use the exact IP address shown in Vercel → Project Settings → Domains when uniware.net is added. "
        "Legacy Vercel IP (if Vercel shows this): 76.76.21.21",
    )
    add_table(
        doc,
        ["Type", "Name / Host", "Value"],
        [
            ["A", "@", "Vercel IP from Domains panel (e.g. 76.76.21.21)"],
        ],
    )
    add_para(doc, "DO NOT modify MX, SPF, or DKIM records.")
    add_para(doc, "Test: open https://uniware.net — should show the new Vercel site.")

    add_heading(doc, "Step 3 — Email smoke test", level=2)
    add_bullet(doc, "Send and receive a test email to sales@uniware.net")
    add_bullet(doc, "Confirm delivery works as before")

    add_heading(doc, "6. After Go-Live — Srimathi (same day)", level=1)
    for item in [
        "Update Vercel env: WORDPRESS_FALLBACK_ORIGIN = direct AWS IP/hostname from Dhana (NOT uniware.net — avoids proxy loop)",
        "Redeploy Vercel project",
        "Verify unbuilt pages still work (e.g. /about-us/, /career/)",
        "Verify key redirects (e.g. /aws/ → new AWS page, /sentinel-one-cloud-security/ → Cloud Security page)",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Post Go-Live Monitoring (48 hours)", level=1)
    for item in [
        "Check Google Search Console → Coverage report daily for 404 errors",
        "Fix any missing redirect rules quickly",
        "Monitor contact form submissions",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Pages Handled Automatically (No Redirect Needed)", level=1)
    add_para(
        doc,
        "These old URLs are NOT redirected. The fallback rule serves them from WordPress until rebuilt:",
    )
    for item in [
        "/about-us/, /career/, /contact-us/ — Company section not yet built",
        "/hybrid-cloud-infrastructure/, /multi-cloud-management-services/ — Multi-cloud not built",
        "/dell-emc-dealer.html, VMware partner pages — Partners section not built",
        "Homepage (/) — new homepage already on Vercel; swaps automatically at cutover",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9. Contacts", level=1)
    add_table(
        doc,
        ["Role", "Name", "Responsibility"],
        [
            ["DNS / Infrastructure", "Dhana", "AWS origin, Bluehost DNS, email verification"],
            ["Sign-off", "Niki", "Approve go-live after testing"],
            ["Development / Vercel", "Srimathi", "Code, env vars, testing, post-cutover updates"],
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Document version: September 2026 · Uniware Systems")
    fr.font.size = Pt(9)
    fr.font.italic = True
    fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
